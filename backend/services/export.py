import io
from typing import List, Dict, Any
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import markdown
from xhtml2pdf import pisa

def format_time(seconds: float) -> str:
    """Format seconds into MM:SS or HH:MM:SS"""
    seconds = int(seconds)
    m, s = divmod(seconds, 60)
    h, m = divmod(m, 60)
    if h > 0:
        return f"{h:02d}:{m:02d}:{s:02d}"
    return f"{m:02d}:{s:02d}"

def export_transcript_docx(transcript_content: List[Dict[str, Any]], filename: str) -> io.BytesIO:
    document = Document()
    document.add_heading(filename, 0)

    for segment in transcript_content:
        start_time = format_time(segment.get("start", 0))
        speaker = segment.get("speaker_id", "Unknown")
        text = segment.get("text", "")

        # Header: [00:00] Speaker
        p = document.add_paragraph()
        runner = p.add_run(f"[{start_time}] {speaker}")
        runner.bold = True
        
        # Text
        document.add_paragraph(text)
        document.add_paragraph("") # Spacer

    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    return file_stream

def export_transcript_pdf(transcript_content: List[Dict[str, Any]], filename: str) -> io.BytesIO:
    # Build HTML
    html = f"<html><body><h1>{filename}</h1>"
    for segment in transcript_content:
        start_time = format_time(segment.get("start", 0))
        speaker = segment.get("speaker_id", "Unknown")
        text = segment.get("text", "")
        
        html += f"<div><p><strong>[{start_time}] {speaker}</strong></p><p>{text}</p><br/></div>"
    html += "</body></html>"

    return _generate_pdf_from_html(html)

def export_minutes_docx(minutes_content: str, filename: str) -> io.BytesIO:
    document = Document()
    document.add_heading(filename, 0)

    # Simple Markdown parser for Docx
    # Split by lines
    lines = minutes_content.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            document.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            document.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            document.add_heading(line[4:], level=3)
        elif line.startswith('- ') or line.startswith('* '):
            document.add_paragraph(line[2:], style='List Bullet')
        else:
            # Handle bold text **text** roughly
            p = document.add_paragraph()
            parts = line.split('**')
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 == 1: # Odd parts are between ** **
                    run.bold = True

    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    return file_stream

def export_minutes_pdf(minutes_content: str, filename: str) -> io.BytesIO:
    # Convert Markdown to HTML
    html_content = markdown.markdown(minutes_content)
    
    full_html = f"""
    <html>
    <head>
        <style>
            body {{ font-family: sans-serif; }}
            h1, h2, h3 {{ color: #333; }}
            ul {{ margin-bottom: 10px; }}
            p {{ margin-bottom: 10px; line-height: 1.5; }}
        </style>
    </head>
    <body>
        <h1>{filename}</h1>
        {html_content}
    </body>
    </html>
    """
    
    return _generate_pdf_from_html(full_html)

def _generate_pdf_from_html(html_content: str) -> io.BytesIO:
    file_stream = io.BytesIO()
    
    # Configure font support for Chinese if needed
    # Note: xhtml2pdf requires a font that supports the characters.
    # Without a specific font file, CJK might not render correctly on all systems.
    # We use a default configuration here.
    
    pisa_status = pisa.CreatePDF(
        src=html_content,
        dest=file_stream,
        encoding='utf-8'
    )
    
    if pisa_status.err:
        raise Exception("PDF generation failed")
        
    file_stream.seek(0)
    return file_stream
